from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel

from config import APP_NAME, FRONTEND_ORIGIN, OUTPUT_DIR
from services.audio_utils import make_output_filename, save_upload_file
from services.stt_service import get_stt_status, transcribe_audio as run_stt
from services.summarizer_service import get_summarizer_status, note_to_markdown, summarize_text
from services.tts_service import get_tts_status, synthesize_speech

app = FastAPI(title=APP_NAME, version="1.2.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[FRONTEND_ORIGIN, "http://localhost:5173", "http://127.0.0.1:5173", "*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.mount("/outputs", StaticFiles(directory=str(OUTPUT_DIR)), name="outputs")


class TextRequest(BaseModel):
    text: str
    summary_mode: str | None = "singkat"
    note_title: str | None = ""
    course_name: str | None = ""


class ExportDocxRequest(BaseModel):
    note_title: str = "Catatan Kuliah Otomatis"
    course_name: str = ""
    summary_mode: str = "singkat"
    transcription: str = ""
    summary: str = ""
    structured_note: dict[str, Any] | None = None
    audio_url: str = ""


class ProcessAllResponse(BaseModel):
    transcription: str
    summary: str
    structured_note: dict[str, Any] | None = None
    audio_url: str
    stt_mode: str
    summarize_mode: str
    summary_mode: str
    tts_mode: str
    note_title: str = ""
    course_name: str = ""
    warnings: list[str] = []


@app.get("/")
def root():
    return {
        "message": APP_NAME,
        "docs": "http://localhost:8000/docs",
        "health": "http://localhost:8000/api/health",
    }


@app.get("/api/health")
def health():
    return {
        "status": "ok",
        "app": APP_NAME,
        "version": "1.2.0-stage2",
        "stt": get_stt_status(),
        "summarizer": get_summarizer_status(),
        "tts": get_tts_status(),
    }


@app.post("/api/transcribe")
async def transcribe_endpoint(
    file: UploadFile = File(...),
    fallback_text: str | None = Form(default=None),
):
    try:
        audio_path = await save_upload_file(file)
        result = run_stt(audio_path, fallback_text=fallback_text)
        return {
            "transcription": result.get("text", ""),
            "text": result.get("text", ""),
            "mode": result.get("mode", "-"),
            "error": result.get("error"),
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/summarize")
def summarize(request: TextRequest):
    try:
        result = summarize_text(
            request.text,
            summary_mode=request.summary_mode or "singkat",
            note_title=request.note_title or "",
            course_name=request.course_name or "",
        )
        return result
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/tts")
def tts(request: TextRequest):
    try:
        result = synthesize_speech(request.text)
        audio_url = f"http://localhost:8000/outputs/{result['filename']}"
        return {
            "audio_url": audio_url,
            "filename": result["filename"],
            "mode": result["mode"],
            "error": result["error"],
        }
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


@app.post("/api/process-all", response_model=ProcessAllResponse)
async def process_all(
    file: UploadFile = File(...),
    fallback_text: str | None = Form(default=None),
    summary_mode: str | None = Form(default="singkat"),
    note_title: str | None = Form(default=""),
    course_name: str | None = Form(default=""),
):
    warnings: list[str] = []
    try:
        audio_path = await save_upload_file(file)

        stt_result = run_stt(audio_path, fallback_text=fallback_text)
        if stt_result.get("error"):
            warnings.append(f"STT: {stt_result['error']}")

        summary_result = summarize_text(
            stt_result["text"],
            summary_mode=summary_mode or "singkat",
            note_title=note_title or "",
            course_name=course_name or "",
        )
        if summary_result.get("error"):
            warnings.append(f"Ringkasan: {summary_result['error']}")

        tts_text = ""
        structured = summary_result.get("structured_note") or {}
        if structured:
            tts_text = str(structured.get("summary") or summary_result["summary"])
        else:
            tts_text = summary_result["summary"]

        tts_result = synthesize_speech(tts_text)
        if tts_result.get("error"):
            warnings.append(f"TTS: {tts_result['error']}")

        audio_url = f"http://localhost:8000/outputs/{tts_result['filename']}"
        return ProcessAllResponse(
            transcription=stt_result["text"],
            summary=summary_result["summary"],
            structured_note=summary_result.get("structured_note"),
            audio_url=audio_url,
            stt_mode=stt_result["mode"],
            summarize_mode=summary_result["mode"],
            summary_mode=summary_result.get("summary_mode", summary_mode or "singkat"),
            tts_mode=tts_result["mode"],
            note_title=note_title or "",
            course_name=course_name or "",
            warnings=warnings,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc


def _add_markdown_like_text(document, text: str) -> None:
    """Menulis ringkasan markdown sederhana ke DOCX tanpa parser tambahan."""
    for raw_line in (text or "").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line[:2].isdigit() and "." in line[:4]:
            document.add_paragraph(line, style="List Number")
        else:
            document.add_paragraph(line)


@app.post("/api/export-docx")
def export_docx(request: ExportDocxRequest):
    try:
        from docx import Document

        filename = make_output_filename("catatan_kuliah", ".docx")
        output_path = OUTPUT_DIR / filename

        document = Document()
        title = request.note_title.strip() or "Catatan Kuliah Otomatis"
        course = request.course_name.strip() or "-"
        document.add_heading(title, level=0)
        document.add_paragraph(f"Mata Kuliah: {course}")
        document.add_paragraph(f"Mode Ringkasan: {request.summary_mode or '-'}")
        document.add_paragraph(f"Tanggal Export: {datetime.now().strftime('%d-%m-%Y %H:%M')}")

        note = request.structured_note or None
        if note:
            document.add_heading("Ringkasan", level=1)
            document.add_paragraph(str(note.get("summary") or "-"))

            document.add_heading("Poin Penting", level=1)
            for item in list(note.get("important_points") or []):
                document.add_paragraph(str(item), style="List Bullet")

            document.add_heading("Kata Kunci", level=1)
            for item in list(note.get("keywords") or []):
                document.add_paragraph(str(item), style="List Bullet")

            document.add_heading("Pertanyaan Latihan", level=1)
            for item in list(note.get("questions") or []):
                document.add_paragraph(str(item), style="List Number")

            document.add_heading("Tindak Lanjut Belajar", level=1)
            for item in list(note.get("action_items") or []):
                document.add_paragraph(str(item), style="List Bullet")

            document.add_heading("Kesimpulan", level=1)
            document.add_paragraph(str(note.get("conclusion") or "-"))
        else:
            _add_markdown_like_text(document, request.summary)

        if request.transcription.strip():
            document.add_heading("Transkripsi", level=1)
            document.add_paragraph(request.transcription.strip())

        if request.audio_url.strip():
            document.add_heading("Audio TTS", level=1)
            document.add_paragraph(request.audio_url.strip())

        document.save(output_path)
        return FileResponse(
            path=str(output_path),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=filename,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=str(exc)) from exc
